import datetime
import json
import os
import uuid
from pathlib import Path
from textwrap import dedent
from xml.sax.saxutils import escape
from zipfile import ZipFile

try:
    from docx import Document  # type: ignore
except ModuleNotFoundError:  # pragma: no cover
    # Minimal fallback implementation for offline environments.
    class _SimpleParagraph:
        def __init__(self, document, text, style):
            self.document = document
            self.text = text
            self.style = style

        def add_run(self, text):
            self.text += text
            return self

        def bold(self):
            return self

    class _TableCell:
        def __init__(self):
            self.text = ""

    class _SimpleTable:
        def __init__(self, rows, cols):
            if rows <= 0 or cols <= 0:
                raise ValueError("rows and cols must be positive")
            self._rows = [[_TableCell() for _ in range(cols)] for _ in range(rows)]

        def cell(self, row, col):
            return self._rows[row][col]

        @property
        def rows(self):
            return self._rows

        @property
        def row_count(self):
            return len(self._rows)

        @property
        def col_count(self):
            return len(self._rows[0]) if self._rows else 0

        def add_row(self):
            new_row = [_TableCell() for _ in range(self.col_count)]
            self._rows.append(new_row)
            return new_row

    class Document:  # type: ignore
        def __init__(self):
            self._elements = []

        def add_paragraph(self, text="", style=None):
            para = _SimpleParagraph(self, text, style)
            self._elements.append(("paragraph", text, style))
            return para

        def add_heading(self, text, level=1):
            style = f"Heading{level}"
            self._elements.append(("paragraph", text, style))

        def add_table(self, rows, cols):
            table = _SimpleTable(rows, cols)
            self._elements.append(("table", table, None))
            return table

        def add_page_break(self):
            self._elements.append(("page_break", None, None))

        def save(self, path):
            body_fragments = []

            def _paragraph_xml(text, style):
                paragraphs = []
                for part in text.split("\n"):
                    if not part:
                        paragraphs.append("<w:p/>")
                        continue
                    style_xml = ""
                    if style:
                        style_xml = f"<w:pPr><w:pStyle w:val=\"{escape(style)}\"/></w:pPr>"
                    paragraphs.append(
                        "<w:p>" + style_xml + f"<w:r><w:t>{escape(part)}</w:t></w:r></w:p>"
                    )
                return "".join(paragraphs)

            def _table_xml(table):
                cols = table.col_count
                col_width = max(1, int(9000 / max(cols, 1)))
                tbl_pr = (
                    "<w:tblPr>"
                    "<w:tblStyle w:val=\"TableGrid\"/>"
                    "<w:tblW w:w=\"0\" w:type=\"auto\"/>"
                    "<w:tblBorders>"
                    "<w:top w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                    "<w:left w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                    "<w:bottom w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                    "<w:right w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                    "<w:insideH w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                    "<w:insideV w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                    "</w:tblBorders>"
                    "</w:tblPr>"
                )
                grid = "".join(
                    f"<w:gridCol w:w=\"{col_width}\"/>" for _ in range(cols)
                )
                rows_xml = []
                for row in table.rows:
                    cells_xml = []
                    for cell in row:
                        text = cell.text or ""
                        paragraph_xml = _paragraph_xml(text, None)
                        cells_xml.append(
                            "<w:tc><w:tcPr><w:tcW w:w=\"0\" w:type=\"auto\"/></w:tcPr>"
                            f"{paragraph_xml}</w:tc>"
                        )
                    rows_xml.append("<w:tr>" + "".join(cells_xml) + "</w:tr>")
                return f"<w:tbl>{tbl_pr}<w:tblGrid>{grid}</w:tblGrid>{''.join(rows_xml)}</w:tbl>"

            for kind, payload, style in self._elements:
                if kind == "paragraph":
                    body_fragments.append(_paragraph_xml(payload, style))
                elif kind == "table":
                    body_fragments.append(_table_xml(payload))
                elif kind == "page_break":
                    body_fragments.append("<w:p><w:r><w:br w:type=\"page\"/></w:r></w:p>")

            sect_pr = (
                "<w:sectPr>"
                "<w:pgSz w:w=12240 w:h=15840/>"
                "<w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\" w:header=\"720\" w:footer=\"720\" w:gutter=\"0\"/>"
                "</w:sectPr>"
            )
            body_xml = "".join(body_fragments) + sect_pr
            document_xml = (
                "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
                "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
                "xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\" "
                "xmlns:w14=\"http://schemas.microsoft.com/office/word/2010/wordml\">"
                f"<w:body>{body_xml}</w:body>"
                "</w:document>"
            )

            styles_xml = dedent(
                """
                <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
                <w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
                    <w:name w:val="Normal"/>
                    <w:basedOn w:val="Normal"/>
                    <w:uiPriority w:val="0"/>
                    <w:qFormat/>
                  </w:style>
                  <w:style w:type="paragraph" w:styleId="Heading1">
                    <w:name w:val="heading 1"/>
                    <w:basedOn w:val="Normal"/>
                    <w:next w:val="Normal"/>
                    <w:uiPriority w:val="9"/>
                    <w:qFormat/>
                    <w:pPr>
                      <w:keepNext/>
                      <w:keepLines/>
                      <w:spacing w:after="200"/>
                    </w:pPr>
                    <w:rPr>
                      <w:b/>
                      <w:sz w:val="32"/>
                    </w:rPr>
                  </w:style>
                  <w:style w:type="paragraph" w:styleId="Heading2">
                    <w:name w:val="heading 2"/>
                    <w:basedOn w:val="Normal"/>
                    <w:next w:val="Normal"/>
                    <w:uiPriority w:val="9"/>
                    <w:qFormat/>
                    <w:rPr>
                      <w:b/>
                      <w:sz w:val="28"/>
                    </w:rPr>
                  </w:style>
                  <w:style w:type="paragraph" w:styleId="Heading3">
                    <w:name w:val="heading 3"/>
                    <w:basedOn w:val="Normal"/>
                    <w:next w:val="Normal"/>
                    <w:uiPriority w:val="9"/>
                    <w:qFormat/>
                    <w:rPr>
                      <w:b/>
                      <w:sz w:val="24"/>
                    </w:rPr>
                  </w:style>
                  <w:style w:type="table" w:styleId="TableGrid">
                    <w:name w:val="Table Grid"/>
                    <w:uiPriority w:val="39"/>
                    <w:rsid/>
                    <w:tblPr>
                      <w:tblBorders>
                        <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                        <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                        <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                      </w:tblBorders>
                    </w:tblPr>
                  </w:style>
                </w:styles>
                """
            ).strip()

            content_types_xml = dedent(
                """
                <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
                <Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
                  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
                  <Default Extension="xml" ContentType="application/xml"/>
                  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
                  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
                  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
                  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
                </Types>
                """
            ).strip()

            rels_xml = dedent(
                """
                <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
                <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
                  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
                  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
                  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
                </Relationships>
                """
            ).strip()

            app_xml = dedent(
                """
                <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
                <Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
                  <Application>python-docx (minimal)</Application>
                  <DocSecurity>0</DocSecurity>
                  <ScaleCrop>false</ScaleCrop>
                  <HeadingPairs>
                    <vt:vector size="2" baseType="variant">
                      <vt:variant><vt:lpstr>Title</vt:lpstr></vt:variant>
                      <vt:variant><vt:i4>1</vt:i4></vt:variant>
                    </vt:vector>
                  </HeadingPairs>
                  <TitlesOfParts>
                    <vt:vector size="1" baseType="lpstr">
                      <vt:lpstr>Ferrisys Document</vt:lpstr>
                    </vt:vector>
                  </TitlesOfParts>
                  <Company>Ferrisys</Company>
                  <LinksUpToDate>false</LinksUpToDate>
                  <SharedDoc>false</SharedDoc>
                  <HyperlinksChanged>false</HyperlinksChanged>
                  <AppVersion>16.0000</AppVersion>
                </Properties>
                """
            ).strip()

            created = datetime.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
            core_xml = dedent(
                f"""
                <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
                <cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
                  <dc:title>Ferrisys Document</dc:title>
                  <dc:creator>Ferrisys Generator</dc:creator>
                  <cp:lastModifiedBy>Ferrisys Generator</cp:lastModifiedBy>
                  <dcterms:created xsi:type="dcterms:W3CDTF">{created}</dcterms:created>
                  <dcterms:modified xsi:type="dcterms:W3CDTF">{created}</dcterms:modified>
                </cp:coreProperties>
                """
            ).strip()

            settings_xml = dedent(
                """
                <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
                <w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                  <w:zoom w:percent="100"/>
                  <w:defaultTabStop w:val="720"/>
                </w:settings>
                """
            ).strip()

            Path(path).parent.mkdir(parents=True, exist_ok=True)
            with ZipFile(path, "w") as zf:
                zf.writestr("[Content_Types].xml", content_types_xml)
                zf.writestr("_rels/.rels", rels_xml)
                zf.writestr("docProps/app.xml", app_xml)
                zf.writestr("docProps/core.xml", core_xml)
                zf.writestr("word/document.xml", document_xml)
                zf.writestr("word/styles.xml", styles_xml)
                zf.writestr("word/settings.xml", settings_xml)


def main() -> None:
    base_dir = Path(__file__).resolve().parents[1]
    docs_dir = base_dir / "docs"
    migration_dir = base_dir / "back-costa" / "src" / "main" / "resources" / "db" / "migration"
    docs_dir.mkdir(parents=True, exist_ok=True)
    migration_dir.mkdir(parents=True, exist_ok=True)

    today = datetime.date.today().strftime("%d/%m/%Y")

    modules = [
        {
            "code": "CORE_AUTH",
            "name": "Core de Autenticación",
            "description": "Gestión centralizada de identidades, roles y permisos con JWT y auditoría.",
            "objective": "Garantizar acceso seguro y trazable a todos los módulos Ferrisys.",
            "functionalities": [
                "Autenticación JWT multi-tenant",
                "Gestión de usuarios, roles y políticas",
                "Asignación dinámica de menús y módulos",
                "Auditoría de sesiones y bitácora de accesos",
            ],
            "status": "Desarrollado",
            "dependencies": [],
            "capabilities": [
                "login",
                "roles",
                "permisos",
                "tokens",
            ],
            "feature_flags": ["otp", "sessionTimeout", "globalCors"],
        },
        {
            "code": "ORG_BRANCH",
            "name": "Sucursales y Organizaciones",
            "description": "Estructura jerárquica de compañías, sucursales y bodegas.",
            "objective": "Modelar multi-empresa y delegar permisos por ubicación.",
            "functionalities": [
                "Registro de empresas y sucursales",
                "Asignación de responsables y zonas",
                "Relación con bodegas e inventarios",
            ],
            "status": "En progreso",
            "dependencies": ["CORE_AUTH", "MASTER_DATA"],
            "capabilities": ["multiTenant", "branchControl", "geocoding"],
            "feature_flags": ["multiBranch", "geoTagging"],
        },
        {
            "code": "MASTER_DATA",
            "name": "Datos Maestros",
            "description": "Catálogos transversales (unidades, categorías, impuestos).",
            "objective": "Centralizar catálogos para asegurar consistencia y trazabilidad.",
            "functionalities": [
                "Gestión de catálogos base",
                "Sincronización con módulos operativos",
                "Soporte de localización y divisas",
            ],
            "status": "En progreso",
            "dependencies": ["CORE_AUTH"],
            "capabilities": ["catalogos", "localizacion", "taxes"],
            "feature_flags": ["multiCurrency", "regionalCatalogs"],
        },
        {
            "code": "PRODUCTS",
            "name": "Productos y Servicios",
            "description": "Definición de SKUs, kits, servicios y atributos técnicos.",
            "objective": "Mantener catálogo completo de bienes comercializados.",
            "functionalities": [
                "Registro de productos con atributos",
                "Asociación de listas de precios y unidades",
                "Integración con inventario y compras",
            ],
            "status": "En progreso",
            "dependencies": ["CORE_AUTH", "MASTER_DATA"],
            "capabilities": ["sku", "attributes", "pricingLink"],
            "feature_flags": ["kitSupport", "variantMatrix"],
        },
        {
            "code": "INVENTORY",
            "name": "Inventario",
            "description": "Control de existencias, movimientos, lotes y series.",
            "objective": "Garantizar trazabilidad y disponibilidad de stock.",
            "functionalities": [
                "Movimientos de entrada y salida",
                "Control de lotes y series opcionales",
                "Reconciliación física y auditorías",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "PRODUCTS", "MASTER_DATA"],
            "capabilities": ["stock", "movimientos", "ajustes"],
            "feature_flags": ["multiWarehouse", "lots", "serials"],
        },
        {
            "code": "DOCUMENTS",
            "name": "Gestión Documental",
            "description": "Plantillas y resguardo de documentos transaccionales.",
            "objective": "Centralizar comprobantes y evidencias adjuntas.",
            "functionalities": [
                "Versionado de plantillas",
                "Almacenamiento en nube privada",
                "Metadatos y búsqueda avanzada",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH"],
            "capabilities": ["storage", "templates", "search"],
            "feature_flags": ["s3Storage", "retentionPolicies"],
        },
        {
            "code": "NOTIFICATIONS",
            "name": "Notificaciones",
            "description": "Motor de alertas por correo, SMS y app.",
            "objective": "Informar oportunamente eventos críticos y operativos.",
            "functionalities": [
                "Plantillas multi-canal",
                "Reglas de activación",
                "Integración con workflows",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "WORKFLOWS"],
            "capabilities": ["email", "sms", "inApp"],
            "feature_flags": ["webPush", "whatsapp"],
        },
        {
            "code": "CLIENTS",
            "name": "Clientes",
            "description": "Gestión 360 de clientes y contactos.",
            "objective": "Centralizar información comercial y crediticia de clientes.",
            "functionalities": [
                "Alta y mantenimiento de clientes",
                "Historial de transacciones y cotizaciones",
                "Clasificación y segmentación",
            ],
            "status": "En progreso",
            "dependencies": ["CORE_AUTH", "MASTER_DATA"],
            "capabilities": ["crm", "creditProfile", "segmentation"],
            "feature_flags": ["scoreCards", "preferredPricing"],
        },
        {
            "code": "SUPPLIERS",
            "name": "Proveedores",
            "description": "Directorio de proveedores, contratos y SLA.",
            "objective": "Optimizar compras con visibilidad de condiciones y entregas.",
            "functionalities": [
                "Gestión de proveedores y contactos",
                "Evaluación de desempeño",
                "Integración con órdenes de compra",
            ],
            "status": "En progreso",
            "dependencies": ["CORE_AUTH", "MASTER_DATA"],
            "capabilities": ["suppliers", "rating", "contracts"],
            "feature_flags": ["slaTracking", "integrationEDI"],
        },
        {
            "code": "PRICING",
            "name": "Gestión de Precios",
            "description": "Listas, descuentos y políticas comerciales.",
            "objective": "Administrar precios dinámicos y promociones.",
            "functionalities": [
                "Listas base y escalonadas",
                "Descuentos por segmento",
                "Reglas de aprobación",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "PRODUCTS"],
            "capabilities": ["priceLists", "discounts", "approvals"],
            "feature_flags": ["dynamicPricing", "volumeDiscounts"],
        },
        {
            "code": "SALES",
            "name": "Ventas",
            "description": "Ciclo comercial completo desde cotización a cierre.",
            "objective": "Controlar pipeline de ventas y convertir oportunidades en ingresos.",
            "functionalities": [
                "Órdenes de venta",
                "Integración con cotizaciones y POS",
                "Aplicación de políticas de precio",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "CLIENTS", "PRODUCTS", "INVENTORY", "PRICING"],
            "capabilities": ["orders", "fulfillment", "pricing"],
            "feature_flags": ["partialFulfillment", "creditHold"],
        },
        {
            "code": "PURCHASE",
            "name": "Compras",
            "description": "Gestión de órdenes de compra y recepción.",
            "objective": "Garantizar abastecimiento oportuno y controlado.",
            "functionalities": [
                "Órdenes de compra",
                "Recepción contra inventario",
                "Conciliación con cuentas por pagar",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "SUPPLIERS", "PRODUCTS", "INVENTORY"],
            "capabilities": ["purchaseOrders", "receipts", "threeWayMatch"],
            "feature_flags": ["automaticReorder", "approvalFlows"],
        },
        {
            "code": "POS",
            "name": "Punto de Venta",
            "description": "Interfaz táctil para ventas en mostrador.",
            "objective": "Acelerar ventas retail con integración a inventario en tiempo real.",
            "functionalities": [
                "Tickets rápidos",
                "Pagos mixtos",
                "Cierres de caja",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "SALES", "INVENTORY", "PRICING"],
            "capabilities": ["fastCheckout", "cashControl", "receiptPrint"],
            "feature_flags": ["offlineMode", "tipManagement"],
        },
        {
            "code": "RETURNS",
            "name": "Devoluciones",
            "description": "Flujos de devoluciones de clientes y proveedores.",
            "objective": "Controlar reingresos y notas de crédito asociadas.",
            "functionalities": [
                "Devoluciones de venta",
                "Reintegración a inventario",
                "Notas de crédito automáticas",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "SALES", "INVENTORY"],
            "capabilities": ["returns", "creditMemo", "inspection"],
            "feature_flags": ["reverseLogistics", "qualityHold"],
        },
        {
            "code": "WMS",
            "name": "WMS",
            "description": "Gestión avanzada de almacenes y ubicaciones.",
            "objective": "Optimizar operaciones logísticas y picking.",
            "functionalities": [
                "Picking dirigido",
                "Cross-docking",
                "Inventario cíclico",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "INVENTORY"],
            "capabilities": ["wavePicking", "slotting", "rfid"],
            "feature_flags": ["rfScanning", "laborTracking"],
        },
        {
            "code": "PRODUCTION",
            "name": "Producción",
            "description": "Órdenes de producción y consumo de materiales.",
            "objective": "Coordinar procesos de manufactura ligera.",
            "functionalities": [
                "Listas de materiales",
                "Control de etapas",
                "Consumo automático",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "PRODUCTS", "INVENTORY"],
            "capabilities": ["bom", "workOrders", "backflush"],
            "feature_flags": ["finiteCapacity", "qualityChecks"],
        },
        {
            "code": "SERVICE_ORDERS",
            "name": "Órdenes de Servicio",
            "description": "Planificación y ejecución de servicios técnicos.",
            "objective": "Brindar soporte postventa y servicios de campo.",
            "functionalities": [
                "Agenda y asignación",
                "Control de garantías",
                "Checklists digitales",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "CLIENTS", "PRODUCTS"],
            "capabilities": ["dispatch", "warranty", "mobileChecklist"],
            "feature_flags": ["fieldApp", "slaAlerts"],
        },
        {
            "code": "AR",
            "name": "Cuentas por Cobrar",
            "description": "Seguimiento de cartera, créditos y cobranzas.",
            "objective": "Acelerar recuperación de ingresos con visibilidad total.",
            "functionalities": [
                "Estados de cuenta",
                "Recordatorios automáticos",
                "Integración con ventas y bancos",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "SALES", "CLIENTS"],
            "capabilities": ["aging", "creditLimits", "collection"],
            "feature_flags": ["dunning", "promissoryNotes"],
        },
        {
            "code": "AP",
            "name": "Cuentas por Pagar",
            "description": "Control de obligaciones con proveedores.",
            "objective": "Optimizar pagos y descuentos por pronto pago.",
            "functionalities": [
                "Programación de pagos",
                "Conciliación con compras",
                "Alertas de vencimiento",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "PURCHASE", "SUPPLIERS"],
            "capabilities": ["paymentSchedule", "matching", "cashflow"],
            "feature_flags": ["paymentBatches", "earlyPayment"],
        },
        {
            "code": "ACCOUNTING",
            "name": "Contabilidad",
            "description": "Libro diario, mayor y estados financieros.",
            "objective": "Proveer información financiera integrada.",
            "functionalities": [
                "Integración contable automática",
                "Estados financieros",
                "Reportes fiscales",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "AR", "AP", "BANKS"],
            "capabilities": ["ledger", "financialStatements", "fiscalReports"],
            "feature_flags": ["multiGaap", "consolidation"],
        },
        {
            "code": "BANKS",
            "name": "Bancos",
            "description": "Conciliaciones y movimientos bancarios.",
            "objective": "Sincronizar movimientos bancarios y flujo de caja.",
            "functionalities": [
                "Conciliación bancaria",
                "Importación de extractos",
                "Control de cuentas",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH"],
            "capabilities": ["bankFeeds", "reconciliation", "cashPosition"],
            "feature_flags": ["openBanking", "cashPooling"],
        },
        {
            "code": "REPORTING_BI",
            "name": "Reportes y BI",
            "description": "Paneles, KPIs y exportaciones.",
            "objective": "Entregar información accionable a directivos y operativos.",
            "functionalities": [
                "Dashboards operativos",
                "Exportaciones XLS/PDF",
                "Conectores BI",
            ],
            "status": "En progreso",
            "dependencies": ["CORE_AUTH"],
            "capabilities": ["dashboards", "exports", "adHoc"],
            "feature_flags": ["embeddedBI", "scheduledReports"],
        },
        {
            "code": "AUDIT_LOGS",
            "name": "Auditoría y Logs",
            "description": "Registro de eventos y cambios críticos.",
            "objective": "Cumplir normativas y brindar trazabilidad completa.",
            "functionalities": [
                "Bitácora centralizada",
                "Alertas de seguridad",
                "Retención configurable",
            ],
            "status": "En progreso",
            "dependencies": ["CORE_AUTH"],
            "capabilities": ["changeLog", "securityAlerts", "retention"],
            "feature_flags": ["siemIntegration", "anomalyDetection"],
        },
        {
            "code": "WORKFLOWS",
            "name": "Workflows",
            "description": "Orquestación de procesos y aprobaciones.",
            "objective": "Automatizar procesos críticos con reglas configurables.",
            "functionalities": [
                "Diseñador de flujos",
                "Aprobaciones multinivel",
                "Integración con notificaciones",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH"],
            "capabilities": ["processDesigner", "approvals", "sla"],
            "feature_flags": ["bpmn", "conditionalRouting"],
        },
        {
            "code": "INTEGRATIONS",
            "name": "Integraciones",
            "description": "Conectores con sistemas externos y marketplace.",
            "objective": "Extender Ferrisys hacia ecosistemas ERP, eCommerce y SAT.",
            "functionalities": [
                "API gateway",
                "Webhooks bidireccionales",
                "Conectores pre-construidos",
            ],
            "status": "Planificado",
            "dependencies": ["CORE_AUTH", "WORKFLOWS"],
            "capabilities": ["apiGateway", "webhooks", "connectors"],
            "feature_flags": ["edi", "satFiscal"],
        },
        {
            "code": "SETTINGS",
            "name": "Configuración",
            "description": "Preferencias globales, personalización y parámetros.",
            "objective": "Configurar reglas operativas y parametrizar módulos.",
            "functionalities": [
                "Feature flags",
                "Parámetros fiscales",
                "Branding por tenant",
            ],
            "status": "En progreso",
            "dependencies": ["CORE_AUTH", "MASTER_DATA"],
            "capabilities": ["featureFlags", "branding", "taxConfig"],
            "feature_flags": ["darkMode", "customDomains"],
        },
    ]

    sprint_schedule = [
        {
            "name": "Sprint 1",
            "start": "20/10/2025",
            "end": "26/10/2025",
            "focus": "Setup base, CI/CD inicial, CORS global, JWT y semillas core.",
            "progress": "12%",
        },
        {
            "name": "Sprint 2",
            "start": "27/10/2025",
            "end": "02/11/2025",
            "focus": "Autenticación, roles, permisos dinámicos y sidebar según perfiles.",
            "progress": "24%",
        },
        {
            "name": "Sprint 3",
            "start": "03/11/2025",
            "end": "09/11/2025",
            "focus": "CRUD Clientes/Proveedores con validaciones, búsqueda y paginación.",
            "progress": "36%",
        },
        {
            "name": "Sprint 4",
            "start": "10/11/2025",
            "end": "16/11/2025",
            "focus": "Cotizaciones con flujo de aprobación degradado (sin PRICING avanzado).",
            "progress": "48%",
        },
        {
            "name": "Sprint 5",
            "start": "17/11/2025",
            "end": "23/11/2025",
            "focus": "Compras enlazadas a inventario, recepción y ajustes básicos.",
            "progress": "60%",
        },
        {
            "name": "Sprint 6",
            "start": "24/11/2025",
            "end": "30/11/2025",
            "focus": "Inventario con movimientos, lotes/series por feature flags.",
            "progress": "72%",
        },
        {
            "name": "Sprint 7",
            "start": "01/12/2025",
            "end": "07/12/2025",
            "focus": "Reportes, auditoría, exportaciones y BI básico.",
            "progress": "86%",
        },
        {
            "name": "Sprint 8",
            "start": "08/12/2025",
            "end": "14/12/2025",
            "focus": "Integraciones mínimas, settings avanzados y hardening final.",
            "progress": "98%",
        },
        {
            "name": "Buffer QA",
            "start": "15/12/2025",
            "end": "20/12/2025",
            "focus": "QA final, hardening, release notes y go-live controlado.",
            "progress": "100%",
        },
    ]

    phases = [
        {
            "phase": "Fase 1",
            "sprints": "Sprint 1",
            "start": "20/10/2025",
            "end": "26/10/2025",
            "description": "Setup base del proyecto, pipelines CI/CD y cimientos de seguridad.",
            "owner": "Líder Técnico Backend",
        },
        {
            "phase": "Fase 2",
            "sprints": "Sprint 2",
            "start": "27/10/2025",
            "end": "02/11/2025",
            "description": "Autenticación, roles, módulos asignados y navegación dinámica.",
            "owner": "Líder Técnico Frontend",
        },
        {
            "phase": "Fase 3",
            "sprints": "Sprint 3",
            "start": "03/11/2025",
            "end": "09/11/2025",
            "description": "CRUD de Clientes y Proveedores con búsquedas y paginación.",
            "owner": "Equipo Fullstack",
        },
        {
            "phase": "Fase 4",
            "sprints": "Sprint 4",
            "start": "10/11/2025",
            "end": "16/11/2025",
            "description": "Cotizaciones y flujo básico de aprobación (degradado si falta Pricing).",
            "owner": "Product Owner",
        },
        {
            "phase": "Fase 5",
            "sprints": "Sprint 5",
            "start": "17/11/2025",
            "end": "23/11/2025",
            "description": "Compras con integración a inventario y recepción controlada.",
            "owner": "Equipo Compras-Operaciones",
        },
        {
            "phase": "Fase 6",
            "sprints": "Sprint 6",
            "start": "24/11/2025",
            "end": "30/11/2025",
            "description": "Inventario con movimientos, lotes y series bajo feature flags.",
            "owner": "Arquitecto de Datos",
        },
        {
            "phase": "Fase 7",
            "sprints": "Sprint 7",
            "start": "01/12/2025",
            "end": "07/12/2025",
            "description": "Reportes, auditoría, exportaciones y BI básico.",
            "owner": "Equipo BI & QA",
        },
        {
            "phase": "Fase 8",
            "sprints": "Sprint 8",
            "start": "08/12/2025",
            "end": "14/12/2025",
            "description": "Integraciones mínimas, settings, hardening y performance pass.",
            "owner": "Equipo DevOps",
        },
        {
            "phase": "Buffer",
            "sprints": "QA / Release",
            "start": "15/12/2025",
            "end": "20/12/2025",
            "description": "Estabilización, QA final, hardening y release controlado.",
            "owner": "Release Manager",
        },
    ]

    def build_module_paragraphs(doc):
        for module in modules:
            doc.add_heading(f"{module['name']} ({module['code']})", level=3)
            doc.add_paragraph(f"Objetivo: {module['objective']}")
            doc.add_paragraph(f"Estado: {module['status']}")
            deps = module["dependencies"]
            doc.add_paragraph(
                "Dependencias técnicas: "
                + (", ".join(deps) if deps else "Sin dependencias directas"),
            )
            doc.add_paragraph("Funcionalidades principales:")
            for item in module["functionalities"]:
                doc.add_paragraph(f"- {item}")

    doc1 = Document()
    doc1.add_heading("Proyecto Ferrisys - Sistema Integral de Gestión para Ferreterías (Qbit-SaaS)", level=1)
    doc1.add_paragraph("Desarrollo de plataforma web con Angular y microservicios en Spring Boot")
    doc1.add_paragraph("Autor: Selvin Santiago Pu Chiguil")
    doc1.add_paragraph(f"Fecha de entrega: {today}")
    doc1.add_page_break()

    doc1.add_heading("Resumen ejecutivo", level=1)
    doc1.add_paragraph(
        "Visión: Automatizar la gestión administrativa, inventarios y finanzas de ferreterías y negocios de materiales."
    )
    doc1.add_paragraph(
        "Stack tecnológico: Angular, Spring Boot, PostgreSQL, Docker, Maven, JWT, Flyway, Lombok, MapStruct."
    )
    doc1.add_paragraph(
        "Beneficios clave: eficiencia operativa, control de inventario, trazabilidad, modularidad, escalabilidad y soporte multitenant."
    )

    doc1.add_heading("Alcance funcional", level=1)
    build_module_paragraphs(doc1)

    doc1.add_heading("Arquitectura del sistema", level=1)
    doc1.add_paragraph(
        "Frontend Angular SPA consume APIs REST autenticadas con JWT emitidos por los microservicios Spring Boot."
    )
    doc1.add_paragraph(
        "Backend: microservicios Spring Boot con Flyway, Lombok, MapStruct, seguridad JWT, auditoría y CORS global."
    )
    doc1.add_paragraph("Base de datos: PostgreSQL con esquemas auditables (created_at/by, updated_at/by).")
    doc1.add_paragraph(
        "Comunicación principal: localhost:4200 ↔ localhost:8081/ferrisys-service vía HTTPS tras reverse proxy."
    )
    doc1.add_paragraph(
        "Diagrama lógico:\n[Angular SPA] → [API Gateway] → [Microservicios Spring Boot] → [PostgreSQL]"
    )

    doc1.add_heading("Cronograma de trabajo", level=1)
    phase_table = doc1.add_table(rows=len(phases) + 1, cols=6)
    headers = ["Fase", "Sprint(s)", "Fecha inicio", "Fecha fin", "Descripción", "Responsable"]
    for idx, header in enumerate(headers):
        phase_table.cell(0, idx).text = header
    for row_idx, phase in enumerate(phases, start=1):
        phase_table.cell(row_idx, 0).text = phase["phase"]
        phase_table.cell(row_idx, 1).text = phase["sprints"]
        phase_table.cell(row_idx, 2).text = phase["start"]
        phase_table.cell(row_idx, 3).text = phase["end"]
        phase_table.cell(row_idx, 4).text = phase["description"]
        phase_table.cell(row_idx, 5).text = phase["owner"]

    doc1.add_paragraph("")
    sprint_table = doc1.add_table(rows=len(sprint_schedule) + 1, cols=4)
    sprint_headers = ["Sprint", "Fecha inicio", "Fecha fin", "Entregables / % avance"]
    for idx, header in enumerate(sprint_headers):
        sprint_table.cell(0, idx).text = header
    for row_idx, sprint in enumerate(sprint_schedule, start=1):
        sprint_table.cell(row_idx, 0).text = sprint["name"]
        sprint_table.cell(row_idx, 1).text = sprint["start"]
        sprint_table.cell(row_idx, 2).text = sprint["end"]
        sprint_table.cell(row_idx, 3).text = f"{sprint['focus']} ({sprint['progress']})"

    doc1.add_heading("Plan de mejoras y extensiones futuras", level=1)
    future_items = [
        "Multitienda/multiempresa completo",
        "Facturación electrónica SAT",
        "Integración IoT (básculas, códigos QR)",
        "Dashboard en tiempo real",
        "Aplicación móvil híbrida",
        "Permisos granulares y políticas avanzadas",
    ]
    for item in future_items:
        doc1.add_paragraph(f"- {item}")

    doc1.add_heading("Conclusiones", level=1)
    doc1.add_paragraph(
        "El proyecto Ferrisys avanza con una base sólida en seguridad y modularidad. Las fases iniciales consolidan autenticación y catálogos, mientras las siguientes habilitan operaciones core, analítica y extensibilidad."
    )
    doc1.add_paragraph(
        "Próximos pasos: reforzar pruebas automatizadas, preparar pilotos con clientes ancla y afinar el roadmap de integraciones estratégicas."
    )
    doc1.add_paragraph(
        "Beneficios inmediatos: visibilidad integral, reducción de errores manuales y plataforma escalable para crecimiento regional."
    )

    doc1_path = docs_dir / "Ferrisys_Cronograma_Proyecto.docx"
    doc1.save(doc1_path)

    markdown_lines = [
        "# Ferrisys - Qbit-SaaS Overview",
        "",
        "## Resumen ejecutivo",
        "- Visión: Automatizar la gestión administrativa, inventarios y finanzas de ferreterías y negocios de materiales.",
        "- Stack: Angular, Spring Boot, PostgreSQL, Docker, Maven, JWT, Flyway, Lombok, MapStruct.",
        "- Beneficios: Eficiencia operativa, control de inventario, trazabilidad, modularidad, escalabilidad, multitenant.",
        "",
        "## Lista de módulos",
    ]

    for module in modules:
        markdown_lines.append(f"### {module['name']} ({module['code']})")
        markdown_lines.append(f"- Objetivo: {module['objective']}")
        markdown_lines.append("- Capacidades: " + ", ".join(module["capabilities"]))
        markdown_lines.append("- Dependencias: " + (", ".join(module["dependencies"]) if module["dependencies"] else "Ninguna"))
        markdown_lines.append("- Feature Flags: " + (", ".join(module["feature_flags"]) if module["feature_flags"] else "N/A"))
        markdown_lines.append(f"- Estado: {module['status']}")
        markdown_lines.append("")

    markdown_lines.append("## Matriz de dependencias")
    markdown_lines.append("")
    markdown_lines.append("| Módulo | Depende de |")
    markdown_lines.append("| --- | --- |")
    for module in modules:
        deps = ", ".join(module["dependencies"]) if module["dependencies"] else "Ninguno"
        markdown_lines.append(f"| {module['code']} | {deps} |")

    markdown_lines.append("")
    markdown_lines.append("## Packs/licenciamiento sugeridos")
    packs = {
        "Starter": ["CORE_AUTH", "CLIENTS", "SUPPLIERS", "PRODUCTS", "INVENTORY", "SALES", "PURCHASE", "REPORTING_BI"],
        "Retail": ["Starter", "POS", "PRICING", "RETURNS"],
        "Distribución": ["Starter", "WMS", "WORKFLOWS"],
        "Manufactura": ["Starter", "PRODUCTION", "WORKFLOWS"],
        "Finanzas": ["Starter", "AR", "AP", "BANKS", "ACCOUNTING"],
        "Enterprise": ["Todo", "INTEGRATIONS", "REPORTING_BI", "AUDIT_LOGS"],
    }
    for name, components in packs.items():
        markdown_lines.append(f"- **{name}**: {', '.join(components)}")
    markdown_lines.append("")

    markdown_lines.append("## Contrato de datos (JSON de ejemplo)")
    data_contract = {
        "customerId": "UUID",
        "branchId": "UUID",
        "documentType": "COTIZACION",
        "items": [
            {"sku": "SKU-001", "quantity": 10, "unitPrice": 120.5, "currency": "GTQ"},
        ],
        "totals": {"subtotal": 1205.0, "tax": 144.6, "grandTotal": 1349.6},
        "metadata": {"approvedBy": "UUID", "approvedAt": "2025-11-20T15:30:00Z"},
    }
    markdown_lines.append("```json")
    markdown_lines.append(json.dumps(data_contract, indent=2, ensure_ascii=False))
    markdown_lines.append("```")
    markdown_lines.append("")

    markdown_lines.append("## Flujos degradados")
    degraded = [
        "Cotizaciones operan sin PRICING avanzado aplicando lista base.",
        "Compras registran recepciones sin INVENTORY avanzado (movimientos simplificados).",
        "POS cae a captura manual si WORKFLOWS o NOTIFICATIONS están inactivos.",
    ]
    for item in degraded:
        markdown_lines.append(f"- {item}")
    markdown_lines.append("")

    markdown_lines.append("## Roadmap de futuras integraciones")
    roadmap = [
        "Conector SAT para facturación electrónica.",
        "Integración con ERPs contables (SAP Business One, Odoo).",
        "API para dispositivos IoT (básculas, sensores de peso).",
        "Marketplace B2B y sincronización con plataformas eCommerce.",
    ]
    for item in roadmap:
        markdown_lines.append(f"- {item}")
    markdown_lines.append("")

    markdown_lines.append("## Tabla de sprints y entregables")
    markdown_lines.append("")
    markdown_lines.append("| Sprint | Fecha inicio | Fecha fin | Entregables | % avance |")
    markdown_lines.append("| --- | --- | --- | --- | --- |")
    for sprint in sprint_schedule:
        markdown_lines.append(
            f"| {sprint['name']} | {sprint['start']} | {sprint['end']} | {sprint['focus']} | {sprint['progress']} |"
        )

    markdown_path = docs_dir / "modules_overview.md"
    markdown_path.write_text("\n".join(markdown_lines), encoding="utf-8")

    doc2 = Document()
    doc2.add_heading("Ferrisys - Módulos y Dependencias", level=1)
    doc2.add_paragraph("Resumen ejecutivo: Ferrisys integra gestión comercial, inventarios y finanzas sobre Angular + Spring Boot + PostgreSQL.")

    doc2.add_heading("Lista de módulos", level=1)
    build_module_paragraphs(doc2)

    doc2.add_heading("Matriz de dependencias", level=1)
    matrix_table = doc2.add_table(rows=len(modules) + 1, cols=2)
    matrix_table.cell(0, 0).text = "Módulo"
    matrix_table.cell(0, 1).text = "Depende de"
    for idx, module in enumerate(modules, start=1):
        matrix_table.cell(idx, 0).text = f"{module['code']}"
        deps = ", ".join(module["dependencies"]) if module["dependencies"] else "Ninguno"
        matrix_table.cell(idx, 1).text = deps

    doc2.add_heading("Packs / licenciamiento sugeridos", level=1)
    for name, components in packs.items():
        doc2.add_paragraph(f"{name}: {', '.join(components)}")

    doc2.add_heading("Contrato de datos (JSON)", level=1)
    json_text = json.dumps(data_contract, indent=2, ensure_ascii=False)
    for line in json_text.split("\n"):
        doc2.add_paragraph(line)

    doc2.add_heading("Flujos degradados", level=1)
    for item in degraded:
        doc2.add_paragraph(f"- {item}")

    doc2.add_heading("Roadmap de integraciones", level=1)
    for item in roadmap:
        doc2.add_paragraph(f"- {item}")

    doc2.add_heading("Tabla de sprints", level=1)
    doc2_table = doc2.add_table(rows=len(sprint_schedule) + 1, cols=5)
    headers_doc2 = ["Sprint", "Fecha inicio", "Fecha fin", "Entregables", "% avance"]
    for idx, header in enumerate(headers_doc2):
        doc2_table.cell(0, idx).text = header
    for row_idx, sprint in enumerate(sprint_schedule, start=1):
        doc2_table.cell(row_idx, 0).text = sprint["name"]
        doc2_table.cell(row_idx, 1).text = sprint["start"]
        doc2_table.cell(row_idx, 2).text = sprint["end"]
        doc2_table.cell(row_idx, 3).text = sprint["focus"]
        doc2_table.cell(row_idx, 4).text = sprint["progress"]

    doc2_path = docs_dir / "Ferrisys_Modulos_Dependencias.docx"
    doc2.save(doc2_path)

    yaml_lines = ["modules:"]
    for module in modules:
        yaml_lines.append(f"  - code: {module['code']}")
        yaml_lines.append(f"    name: '{module['name']}'")
        yaml_lines.append(f"    description: '{module['description']}'")
        depends = module["dependencies"]
        if depends:
            yaml_lines.append("    dependsOn:")
            for dep in depends:
                yaml_lines.append(f"      - {dep}")
        else:
            yaml_lines.append("    dependsOn: []")
        yaml_lines.append("    capabilities:")
        for capability in module["capabilities"]:
            yaml_lines.append(f"      - {capability}")
        yaml_lines.append("    featureFlags:")
        for flag in module["feature_flags"]:
            yaml_lines.append(f"      - {flag}")
        yaml_lines.append("    status: 1")

    yaml_path = docs_dir / "modules_catalog.yaml"
    yaml_path.write_text("\n".join(yaml_lines) + "\n", encoding="utf-8")

    sql_lines = [
        "-- Seed de módulos base Ferrisys",
    ]
    for module in modules:
        module_uuid = uuid.uuid5(uuid.NAMESPACE_URL, f"ferrisys/{module['code']}")
        sql_lines.append(
            "INSERT INTO auth_module (id, name, description, status) VALUES "
            f"('{module_uuid}', '{module['name'].replace("'", "''")}', '{module['description'].replace("'", "''")}', 1);"
        )

    sql_path = migration_dir / "V99__seed_auth_modules.sql"
    sql_path.write_text("\n".join(sql_lines) + "\n", encoding="utf-8")

    print("✅ Documentación Ferrisys generada exitosamente en /docs y /migration.")


if __name__ == "__main__":
    main()
